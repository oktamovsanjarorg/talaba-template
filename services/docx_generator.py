import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_page_number(run):
    """Word hujjati sahifasiga avtomatik sahifa raqami XML elementini kiritish"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def format_academic_paragraph(p, text: str, font_size=14, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, is_heading=False, space_before=0, space_after=0):
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if not is_heading:
        p.paragraph_format.first_line_indent = Cm(1.25)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    return run


def create_referat_docx(data: dict, output_path: str, student_name: str = "Talaba", university: str = "O'zbekiston Milliy Universiteti", faculty: str = "Axborot texnologiyalari fakulteti"):
    """
    O'zbekiston OTM standartlaridagi rasmiy akademik Referat (.docx) generatori.
    """
    doc = Document()

    # SECTION 1: TITUL VARAG'I (Sahifa raqamisiz)
    sec_title = doc.sections[0]
    sec_title.top_margin = Cm(2.0)
    sec_title.bottom_margin = Cm(2.0)
    sec_title.left_margin = Cm(3.0)
    sec_title.right_margin = Cm(1.5)
    sec_title.different_first_page_header_footer = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("O'ZBEKISTON RESPUBLIKASI\nOLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI\n\n")
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{university.upper()}\n{faculty.upper()}\n\n\n")
    r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REFERAT\n\n")
    r.font.name = 'Times New Roman'; r.font.size = Pt(22); r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Mavzu: \"{data.get('title', 'Mavzu')}\"\n\n\n\n")
    r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(f"Bajardi: {student_name}\nTekshirdi: Ilmiy rahbar / O'qituvchi\n\n\n\n\n")
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Toshkent — 2026")
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    # SECTION 2: ASOSIY MATN (Sahifa raqamlari bilan)
    sec_main = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec_main.top_margin = Cm(2.0)
    sec_main.bottom_margin = Cm(2.0)
    sec_main.left_margin = Cm(3.0)
    sec_main.right_margin = Cm(1.5)

    footer = sec_main.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = f_p.add_run()
    f_run.font.name = 'Times New Roman'
    f_run.font.size = Pt(11)
    add_page_number(f_run)

    # 1. MUNDARIJA
    p_plan_title = doc.add_paragraph()
    format_academic_paragraph(p_plan_title, "MUNDARIJA", font_size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, is_heading=True, space_after=6)

    page_counter = 3
    for idx, item in enumerate(data.get('plan', []), 1):
        p_item = doc.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_item.paragraph_format.line_spacing = 1.3
        p_item.paragraph_format.space_after = Pt(2)
        dots_count = max(5, 55 - len(item))
        dots = "." * dots_count
        r_item = p_item.add_run(f"{item} {dots} {page_counter}")
        r_item.font.name = 'Times New Roman'; r_item.font.size = Pt(13)
        page_counter += 2

    doc.add_page_break()

    # 2. KIRISH QISMI
    p_intro_title = doc.add_paragraph()
    format_academic_paragraph(p_intro_title, "KIRISH", font_size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, is_heading=True, space_after=6)

    intro_text = data.get('introduction', '')
    for par in intro_text.split("\n\n"):
        if par.strip():
            p_par = doc.add_paragraph()
            format_academic_paragraph(p_par, par.strip(), font_size=14)

    doc.add_page_break()

    # 3. ASOSIY BOBLAR VA BO'LIMLAR
    for idx, chapter in enumerate(data.get('chapters', []), 1):
        ch_title = chapter.get('title', f"{idx}-BOB")
        p_ch = doc.add_paragraph()
        format_academic_paragraph(p_ch, ch_title.upper(), font_size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, is_heading=True, space_before=6, space_after=8)

        # Agar bo'limlar (sections) strukturasi bo'lsa
        sections = chapter.get('sections', [])
        if sections:
            for s in sections:
                subtitle = s.get('subtitle', '')
                if subtitle:
                    p_sub = doc.add_paragraph()
                    format_academic_paragraph(p_sub, subtitle, font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, is_heading=True, space_before=6, space_after=4)
                text = s.get('text', '')
                for par in text.split("\n\n"):
                    if par.strip():
                        p_par = doc.add_paragraph()
                        format_academic_paragraph(p_par, par.strip(), font_size=14)
        else:
            # Oddiy content matni bo'lsa
            content = chapter.get('content', '')
            for par in content.split("\n\n"):
                if par.strip():
                    p_par = doc.add_paragraph()
                    format_academic_paragraph(p_par, par.strip(), font_size=14)
        
        doc.add_page_break()

    # 4. XULOSA
    p_concl_title = doc.add_paragraph()
    format_academic_paragraph(p_concl_title, "XULOSA VA AMALIY TAVSIYALAR", font_size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, is_heading=True, space_after=6)

    concl_text = data.get('conclusion', '')
    for par in concl_text.split("\n\n"):
        if par.strip():
            p_par = doc.add_paragraph()
            format_academic_paragraph(p_par, par.strip(), font_size=14)

    doc.add_page_break()

    # 5. FOYDALANILGAN ADABIYOTLAR
    p_ref_title = doc.add_paragraph()
    format_academic_paragraph(p_ref_title, "FOYDALANILGAN ADABIYOTLAR RO'YXATI", font_size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, is_heading=True, space_after=6)

    for idx, ref in enumerate(data.get('references', []), 1):
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.line_spacing = 1.3
        p_ref.paragraph_format.space_after = Pt(3)
        clean_ref = ref if ref.startswith(f"{idx}.") else f"{idx}. {ref}"
        r_ref = p_ref.add_run(clean_ref)
        r_ref.font.name = 'Times New Roman'
        r_ref.font.size = Pt(13)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
