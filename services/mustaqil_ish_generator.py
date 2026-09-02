import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_mustaqil_ish_docx(data: dict, output_path: str, student_name: str = "Talaba", university: str = "O'zbekiston Oliy Ta'lim Muassasasi"):
    """
    Mustaqil ish uchun OTM talablaridagi toza Word hujjati.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.6)

    # 1. TITUL VARAG'I
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{university.upper()}\n\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MUSTAQIL TA'LIM ISHI\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(20)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Mavzu: \"{data.get('title', 'Mustaqil ish')}\"\n\n\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Bajardi: {student_name}\nQabul qildi: O'qituvchi\n\n\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Toshkent — {datetime.now().year}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    doc.add_page_break()

    # 2. MUNDARIJA / REJA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MUSTAQIL ISH REJASI:\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(15)
    r.bold = True

    for item in data.get('plan', []):
        p = doc.add_paragraph()
        r = p.add_run(f"• {item}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

    doc.add_page_break()

    def add_sec(title: str, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title.upper())
        r.font.name = 'Times New Roman'
        r.font.size = Pt(15)
        r.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.5)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        doc.add_paragraph()

    add_sec("1. TOPSHIRIQNING MAQSADI VA VAZIFALARI", data.get("goal", ""))
    add_sec("2. NAZARIY QISM", data.get("theoretical_part", ""))
    add_sec("3. AMALIY TAHLIL VA YONDASHUV", data.get("practical_part", ""))
    add_sec("4. XULOSA", data.get("conclusion", ""))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FOYDALANILGAN ADABIYOTLAR")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(15)
    r.bold = True

    for ref in data.get('references', []):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(ref)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
