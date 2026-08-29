import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_quiz_docx(quiz_data: list, topic: str, output_path: str):
    """
    Test savollari va javoblar kalitini o'z ichiga olgan chiroyli Word hujjati.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Sarlavha
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"MAVZU: \"{topic.upper()}\"\nNAZORAT VA TEST SAVOLLARI\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(16)
    r.bold = True

    # Savollar
    letters = ["A)", "B)", "C)", "D)"]
    for idx, q in enumerate(quiz_data, 1):
        p_q = doc.add_paragraph()
        r_q = p_q.add_run(f"{idx}. {q.get('question')}")
        r_q.font.name = 'Times New Roman'
        r_q.font.size = Pt(13)
        r_q.bold = True

        for opt_idx, opt in enumerate(q.get('options', [])):
            p_opt = doc.add_paragraph()
            p_opt.paragraph_format.left_indent = Inches(0.3)
            r_opt = p_opt.add_run(f"{letters[opt_idx]} {opt}")
            r_opt.font.name = 'Times New Roman'
            r_opt.font.size = Pt(12)
        doc.add_paragraph()

    # Javoblar kaliti (Alohida sahifa)
    doc.add_page_break()
    p_ans_title = doc.add_paragraph()
    p_ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ans = p_ans_title.add_run("JAVOBLAR KALITI VA IZOHLAR\n")
    r_ans.font.name = 'Times New Roman'
    r_ans.font.size = Pt(15)
    r_ans.bold = True

    for idx, q in enumerate(quiz_data, 1):
        correct_idx = q.get('correct_index', 0)
        options = q.get('options', [])
        correct_text = options[correct_idx] if correct_idx < len(options) else ""
        
        p_k = doc.add_paragraph()
        r_k = p_k.add_run(f"{idx}-savol javobi: {letters[correct_idx]} {correct_text}\n")
        r_k.font.name = 'Times New Roman'
        r_k.font.size = Pt(12)
        r_k.bold = True

        exp = q.get('explanation')
        if exp:
            r_exp = p_k.add_run(f"Izoh: {exp}\n")
            r_exp.font.name = 'Times New Roman'
            r_exp.font.size = Pt(11)
            r_exp.italic = True

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
